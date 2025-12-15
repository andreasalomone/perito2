import datetime
import io
import logging

import re
from typing import List

from docx import Document
from docx.enum.table import (  # Per l'allineamento verticale nelle celle
    WD_CELL_VERTICAL_ALIGNMENT,
)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ITALIAN_MONTHS = {
    1: "gennaio",
    2: "febbraio",
    3: "marzo",
    4: "aprile",
    5: "maggio",
    6: "giugno",
    7: "luglio",
    8: "agosto",
    9: "settembre",
    10: "ottobre",
    11: "novembre",
    12: "dicembre",
}

from app.core.config import settings




def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """Imposta i margini interni di una cella (in dxa, 1 inch = 1440 dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for m_name, m_val in [
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ]:
        mar_el = OxmlElement(f"w:{m_name}")
        mar_el.set(qn("w:w"), str(m_val))
        mar_el.set(qn("w:type"), "dxa")
        tcMar.append(mar_el)
    tcPr.append(tcMar)


def remove_table_borders(table):
    """Rende invisibili tutti i bordi di una tabella."""
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
                border_el = OxmlElement(f"w:{border_name}")
                border_el.set(qn("w:val"), "nil")  # 'nil' o 'none' per nessun bordo

                tcBorders.append(border_el)
            tcPr.append(tcBorders)


logger = logging.getLogger(__name__)





def create_styled_docx(plain_text_report_content: str) -> io.BytesIO:
    if not isinstance(plain_text_report_content, str):
        raise TypeError(
            f"Expected plain_text_report_content to be a string, "
            f"but got {type(plain_text_report_content)}. "
            f"Please check the caller in app.py to ensure 'report_content' is a string instance."
        )
    document: Document = Document()

    # --- Impostazioni di Sezione (Margini, Header/Footer distance) ---
    section = document.sections[0]
    section.top_margin = Cm(4.08)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    # --- Impostazioni di Stile di Default per l'Intero Documento ---
    style = document.styles["Normal"]
    font = style.font
    font.name = settings.DOCX_FONT_NAME
    font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Giustifica il testo
    paragraph_format.line_spacing = (
        settings.DOCX_LINE_SPACING if hasattr(settings, "DOCX_LINE_SPACING") else 1.5
    )
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(
        settings.DOCX_SPACE_AFTER_PARAGRAPH
        if hasattr(settings, "DOCX_SPACE_AFTER_PARAGRAPH")
        else 0
    )

    # --- Header Removed for Default Template ---
    # No logo in the default template.

    # --- Elaborazione del Contenuto Testuale Generato dall'LLM ---
    lines: List[str] = plain_text_report_content.split("\n")
    subject_line_pattern = re.compile(r"^\s*Oggetto\s*:\s*(.*)", re.IGNORECASE)
    date_line_pattern = re.compile(r"^\s*Genova,.*", re.IGNORECASE)
    reference_line_pattern = re.compile(
        r"^\s*(Vs\. Rif\.|Polizza|Ns\. Rif\.)\s*:\s*(.*)"
    )  # Per i riferimenti tipo Vs. Rif.
    section_title_pattern = re.compile(
        r"^\s*([0-9]+)\s*–\s*([A-Z\sÀ-Ù'&]+)\s*$"
    )  # Es. "1 – DATI GENERALI"

    is_first_content_paragraph_after_initial_blocks = (
        True  # Per gestire spazio prima del primo paragrafo narrativo
    )
    initial_right_aligned_lines_count = 0  # Counter for the first five lines

    # --- State for table processing ---
    is_in_table_block = False
    table_lines = []
    dati_generali_lines = []
    is_in_dati_generali_block = False

    for line_num, line in enumerate(lines):
        stripped_line = line.strip()
        original_line = line

        # --- TABLE BLOCK DETECTION (DANNI) ---
        if "[INIZIO_TABELLA_DANNI]" in stripped_line:
            is_in_table_block = True
            continue
        elif "[FINE_TABELLA_DANNI]" in stripped_line:
            is_in_table_block = False
            # Process the collected table lines
            if table_lines:
                _create_damage_table(document, table_lines)
            table_lines = []  # Reset for next potential table
            continue
        elif is_in_table_block:
            if stripped_line:  # Ignore empty lines within the block
                table_lines.append(original_line)
            continue

        # --- TABLE BLOCK DETECTION (DATI GENERALI) ---
        if "[INIZIO_DATI_GENERALI]" in stripped_line:
            is_in_dati_generali_block = True
            continue
        elif "[FINE_DATI_GENERALI]" in stripped_line:
            is_in_dati_generali_block = False
            if dati_generali_lines:
                _create_dati_generali_table(document, dati_generali_lines)
            dati_generali_lines = []
            continue
        elif is_in_dati_generali_block:
            if stripped_line:  # Ignore empty lines
                dati_generali_lines.append(stripped_line)
            continue

        # --- END TABLE BLOCK DETECTION ---

        # Force the first five lines to have a large left indent, creating the desired visual alignment.
        if initial_right_aligned_lines_count < 5:
            p = document.add_paragraph(stripped_line if stripped_line else " ")

            # Set alignment to LEFT and add a large left_indent, as per the user's reference image.
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt = p.paragraph_format
            fmt.left_indent = Cm(
                10
            )  # From the reference image "Indents: Left: 9.99 cm"

            # Minimal spacing for this block
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 1.25
            if (
                initial_right_aligned_lines_count == 0 and line_num == 0
            ):  # Only for the very first line of the document
                fmt.space_before = Pt(10)  # Add some space after the header/logo

            # Make the second line bold as requested.
            if initial_right_aligned_lines_count == 1:
                for run in p.runs:
                    run.bold = True

            initial_right_aligned_lines_count += 1
            continue

        # 2. Gestione Blocco Data (sempre "Genova, [data]")
        if date_line_pattern.match(stripped_line):
            now = datetime.date.today()
            new_date_text = f"Genova, {now.day} {ITALIAN_MONTHS[now.month]} {now.year}"

            p = document.add_paragraph(new_date_text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt = p.paragraph_format
            fmt.first_line_indent = Pt(0)
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            is_first_content_paragraph_after_initial_blocks = (
                True  # Reimposta per il prossimo blocco
            )
            continue

        # 3. Gestione Riferimenti (Vs. Rif., Polizza, Ns. Rif.)
        ref_match = reference_line_pattern.match(stripped_line)
        if ref_match:
            # Create paragraph with 'List Bullet' style for automatic bullet character.
            p = document.add_paragraph(stripped_line, style="List Bullet")
            fmt = p.paragraph_format

            # Apply formatting based on user images.
            fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(0)
            fmt.line_spacing = 2  # User override from instruction.

            # Indentation based on UI:
            # Text starts at Bullet Indent + Text Indent from margin (0.63 + 0.63 = 1.26cm).
            # The bullet hangs to the left of the text.
            fmt.left_indent = Cm(1.26)
            fmt.first_line_indent = Cm(-0.63)

            # Pagination and Hyphenation settings.
            fmt.widow_control = True

            # Disable hyphenation for this paragraph using OXML.
            pPr = p._p.get_or_add_pPr()
            hyphenation = OxmlElement("w:hyphenation")
            hyphenation.set(qn("w:val"), "0")  # '0' corresponds to 'false'
            pPr.append(hyphenation)

            # Font styling for the runs within the paragraph.
            for run in p.runs:
                run.font.name = settings.DOCX_FONT_NAME
                run.font.size = Pt(12)
                run.bold = True

            if is_first_content_paragraph_after_initial_blocks:
                fmt.space_before = Pt(10)
                is_first_content_paragraph_after_initial_blocks = False
            continue

        # 4. Gestione "Oggetto:" in una tabella (borders visible, autofit)
        subject_match = subject_line_pattern.match(original_line)
        if subject_match:
            full_subject_text = subject_match.group(0).strip()
            label_text = "Oggetto:"
            content_after_oggetto = full_subject_text[len(label_text) :].strip()

            # Calcola la larghezza totale disponibile per la tabella
            content_width = (
                document.sections[0].page_width
                - document.sections[0].left_margin
                - document.sections[0].right_margin
            )

            table = document.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            table.autofit = False
            table.allow_autofit = False

            # Imposta larghezze colonne
            col1_width = Cm(2.43)  # Larghezza fissa per "Oggetto:"
            col2_width = content_width - col1_width
            table.columns[0].width = col1_width
            table.columns[1].width = col2_width

            # --- Cella 1: "Oggetto:" ---
            cell_label = table.cell(0, 0)
            cell_label.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para_label = cell_label.paragraphs[0]
            para_label.text = label_text
            para_label.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for run in para_label.runs:
                run.font.name = settings.DOCX_FONT_NAME
                run.font.size = Pt(
                    settings.DOCX_FONT_SIZE_NORMAL
                )  # Dovrebbe essere 12pt
                run.bold = True

            # --- Cella 2: Contenuto dell'oggetto ---
            cell_content = table.cell(0, 1)
            cell_content.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para_content = cell_content.paragraphs[0]
            para_content.text = (
                ""  # Svuota il paragrafo per aggiungere i run formattati
            )
            para_content.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para_content.paragraph_format.line_spacing = 1.0

            # Regex per trovare il testo da mettere in grassetto tra "Ass.to " e " -"
            pattern = re.compile(r"^(Ass\.to\s+)(.+?)(?=\s*-)", re.IGNORECASE)
            match = pattern.match(content_after_oggetto)

            if match:
                # Aggiungi la parte prima del grassetto
                run1 = para_content.add_run(match.group(1))
                run1.font.name = settings.DOCX_FONT_NAME
                run1.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
                run1.bold = False

                # Aggiungi il nome della ditta in grassetto
                run2 = para_content.add_run(match.group(2))
                run2.font.name = settings.DOCX_FONT_NAME
                run2.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
                run2.bold = True

                # Aggiungi il resto della stringa (dal trattino in poi)
                rest_of_string = content_after_oggetto[match.end() :]
                run3 = para_content.add_run(rest_of_string)
                run3.font.name = settings.DOCX_FONT_NAME
                run3.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
                run3.bold = False
            else:
                # Fallback se il pattern non corrisponde: aggiungi tutto il testo non in grassetto
                run = para_content.add_run(content_after_oggetto)
                run.font.name = settings.DOCX_FONT_NAME
                run.font.size = Pt(settings.DOCX_FONT_SIZE_NORMAL)
                run.bold = False

            # Aggiungi uno spazio *dopo* la tabella dell'oggetto
            p_after_table = document.add_paragraph()
            p_after_table.paragraph_format.space_before = Pt(
                0
            )  # Nessuno spazio prima se quello normale è 6pt
            p_after_table.paragraph_format.space_after = Pt(
                settings.DOCX_SPACE_AFTER_PARAGRAPH
                if hasattr(settings, "DOCX_SPACE_AFTER_PARAGRAPH")
                and settings.DOCX_SPACE_AFTER_PARAGRAPH > 0
                else 6
            )  # Spazio dopo l'oggetto
            is_first_content_paragraph_after_initial_blocks = True
            continue

        # 4.5 Gestione "A seguito del gradito incarico"
        if stripped_line.startswith("A seguito del gradito incarico"):
            p = document.add_paragraph(stripped_line)
            p.paragraph_format.first_line_indent = Cm(1.25)
            is_first_content_paragraph_after_initial_blocks = False
            continue

        # 5. Gestione Titoli di Sezione (es. "1 – DATI GENERALI")
        section_title_match = section_title_pattern.match(stripped_line)
        if section_title_match:
            p = document.add_paragraph(stripped_line)
            fmt = p.paragraph_format
            fmt.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
            )  # Assicura che i titoli rimangano allineati a sinistra

            # Applica lo stile grassetto e corsivo ai titoli di sezione come richiesto.
            for run in p.runs:
                run.bold = True
                run.italic = True  # FIX: Aggiungi corsivo

            fmt.space_before = Pt(
                12 if not is_first_content_paragraph_after_initial_blocks else 6
            )  # Più spazio prima dei titoli di sezione
            fmt.space_after = Pt(6)  # Spazio dopo il titolo
            is_first_content_paragraph_after_initial_blocks = False
            continue

        # Handle final disclaimer lines
        if stripped_line.startswith(
            "Il presente certificato di perizia viene emesso"
        ) or stripped_line.startswith("Gli scriventi si riservano il diritto"):
            p = document.add_paragraph(stripped_line)
            p.paragraph_format.first_line_indent = Cm(
                1.25
            )  # Ripristina l'indentazione di default
            p.paragraph_format.line_spacing = (
                1.0  # Override del line spacing di default
            )
            for run in p.runs:
                run.italic = True
            continue

        # 6. Gestione Paragrafi Normali, Liste e Contenuto Indentato
        if not stripped_line and line.strip() == "":  # Riga vuota per spaziatura
            if line_num > 0 and lines[line_num - 1].strip():
                document.add_paragraph()  # Crea lo spazio dato da \n\n (usa space_after di default)
            is_first_content_paragraph_after_initial_blocks = (
                True  # Potrebbe essere una separazione prima di un nuovo blocco
            )
            # continue - redundant at end of loop
        elif stripped_line:
            p = document.add_paragraph()
            fmt = p.paragraph_format

            # Paragrafo standard
            p.text = stripped_line
            # Apply a first-line indent to all standard paragraphs
            fmt.first_line_indent = Cm(1.25)
            if (
                is_first_content_paragraph_after_initial_blocks
                and not section_title_match
            ):
                # Aggiungi spazio prima del primo paragrafo effettivo
                prev_line_empty_or_subject = (
                    line_num > 0 and not lines[line_num - 1].strip()
                ) or (
                    line_num > 0
                    and subject_line_pattern.match(lines[line_num - 1].strip())
                )
                if prev_line_empty_or_subject:
                    fmt.space_before = Pt(
                        0
                    )  # Spazio prima del primo paragrafo effettivo
                    fmt.space_after = Pt(0)
            is_first_content_paragraph_after_initial_blocks = False
            # continue - redundant at end of loop

    # --- Footer con layout Semplificato: Numero Pagina in alto a dx ---
    section = document.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False  # Assicura footer specifico per questa sezione

    # Svuota il footer da paragrafi preesistenti per sicurezza
    for para in footer.paragraphs:
        para._p.getparent().remove(para._p)

    # 1. Paragrafo per il numero di pagina
    p_page_num = footer.add_paragraph()
    p_page_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Rimuovi qualsiasi indentazione ereditata dallo stile 'Normal'
    fmt_page_num = p_page_num.paragraph_format
    fmt_page_num.first_line_indent = Pt(0)
    fmt_page_num.space_before = Pt(0)
    fmt_page_num.space_after = Pt(0)
    fmt_page_num.line_spacing = 1.0

    run_page = p_page_num.add_run()
    run_page.font.name = settings.DOCX_FONT_NAME
    run_page.font.size = Pt(12)

    # Aggiungi campo PAGE per il numero di pagina
    fldChar_begin = OxmlElement("w:fldChar")
    fldChar_begin.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar_end = OxmlElement("w:fldChar")
    fldChar_end.set(qn("w:fldCharType"), "end")

    run_page._r.append(fldChar_begin)
    run_page._r.append(instrText)
    run_page._r.append(fldChar_end)

    # NO Static Footer Text for Default Template

    # Salva in un oggetto BytesIO
    file_stream: io.BytesIO = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream


def _create_damage_table(document, lines: List[str]):
    """Creates and styles a real DOCX table from text lines."""
    if not lines:
        return

    # Determine column count from the header row (first line)
    header_cols = re.split(r"\s{2,}", lines[0].strip())
    num_cols = len(header_cols)
    if num_cols == 0:
        return

    table = document.add_table(rows=0, cols=num_cols)
    table.autofit = True
    table.style = "Table Grid"

    for i, line in enumerate(lines):
        stripped_line = line.strip()
        if not stripped_line:
            continue

        row_cells = table.add_row().cells
        # Split the row by 2 or more spaces to get columns
        cols = re.split(r"\s{2,}", stripped_line)

        # Distribute columns into cells, checking for the 'wxyz' placeholder
        for j, cell_content in enumerate(cols):
            if j < len(row_cells):
                # If the content is the placeholder, insert an empty string.
                # Otherwise, insert the content.
                if cell_content.strip().lower() == "wxyz":
                    row_cells[j].text = ""
                else:
                    row_cells[j].text = cell_content

        # If this is the header row (the first line), make its content bold.
        if i == 0:
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

    # Apply formatting to all cells in the table
    for row in table.rows:
        for cell in row.cells:
            # Set vertical alignment to top
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            # Format paragraph inside the cell
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                fmt = p.paragraph_format
                fmt.line_spacing = 1.0
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)

    # Add some space after the table
    document.add_paragraph()


def _create_dati_generali_table(document, lines: List[str]):
    """Creates a 3-column table for Dati Generali with visible borders."""

    # Calcola larghezze colonne
    content_width = (
        document.sections[0].page_width
        - document.sections[0].left_margin
        - document.sections[0].right_margin
    )
    col1_width = Cm(6.81)
    col2_width = Cm(0.49)
    col3_width = content_width - col1_width - col2_width

    table = document.add_table(rows=0, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    table.allow_autofit = False

    table.columns[0].width = col1_width
    table.columns[1].width = col2_width
    table.columns[2].width = col3_width

    cell_inset_dxa = int((4 / 72) * 1440)  # 4pt Text Inset

    # --- New, more robust logic for processing table data ---
    # This logic groups multi-line values together before creating table rows.

    data_to_render = []
    current_label = None
    current_value_lines = []

    # First pass: Group lines by label
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue

        if ":" in stripped_line:
            # When a new label is found, save the previous one (if it exists)
            if current_label is not None:
                full_value = "\n".join(current_value_lines)
                data_to_render.append({"label": current_label, "value": full_value})

            # Start a new entry
            label, value = stripped_line.split(":", 1)
            current_label = label.strip()
            # Handle cases where the value is on the same line or the next line
            if value.strip():
                current_value_lines = [value.strip()]
            else:
                current_value_lines = []
        else:
            # This is a continuation line, add it to the current value
            if current_label is not None:
                current_value_lines.append(stripped_line)

    # Add the very last entry after the loop finishes
    if current_label is not None:
        full_value = "\n".join(current_value_lines)
        data_to_render.append({"label": current_label, "value": full_value})

    # Second pass: Render the grouped data into the table
    for item in data_to_render:
        row_cells = table.add_row().cells

        # Column 1: Label
        row_cells[0].text = item["label"]

        # Column 2: Colon
        row_cells[1].text = ":"

        # Column 3: Value (potentially multi-line)
        row_cells[2].text = item["value"]

    # Apply formatting to all cells
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            set_cell_margins(
                cell,
                top=cell_inset_dxa,
                start=cell_inset_dxa,
                bottom=cell_inset_dxa,
                end=cell_inset_dxa,
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

            for p in cell.paragraphs:
                # Apply bold to all text
                for run in p.runs:
                    run.bold = True

                # Set alignment
                if i in [1, 2]:  # Colon and Content
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:  # Label
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # Set line spacing
                fmt = p.paragraph_format
                fmt.line_spacing = 1.0
                fmt.space_before = Pt(0)
                fmt.space_after = Pt(0)

    # Add some space after the table
    document.add_paragraph()
