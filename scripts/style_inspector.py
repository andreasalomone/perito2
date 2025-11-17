import argparse
import sys

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def get_font_info(font):
    """Helper function to safely retrieve font information."""
    if not font:
        return "N/A"

    size_pt = "N/A"
    if font.size:
        size_pt = f"{font.size.pt:.2f}pt"

    return (
        f"Font: {font.name or 'Default'}, "
        f"Size: {size_pt}, "
        f"Bold: {font.bold}, Italic: {font.italic}, Underline: {font.underline}, "
        f"Color: {font.color.rgb if font.color and font.color.rgb else 'Auto'}"
    )


def get_cell_border_info(cell):
    """Helper to extract border information from a cell's XML properties."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        return "Borders: Default"

    border_info = []
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = tcBorders.find(qn(f"w:{border_name}"))
        if border_el is not None:
            val = border_el.get(qn("w:val"))
            size = border_el.get(qn("w:sz"))
            color = border_el.get(qn("w:color"))
            info = f"{border_name}: {val}"
            if size:
                info += f" (size: {size})"
            if color:
                info += f" (color: {color})"
            border_info.append(info)

    return " | ".join(border_info) if border_info else "Borders: None"


def inspect_docx_styles(file_path: str):
    """
    Inspects a .docx file and prints its section and style formatting to the console.
    """
    try:
        document = Document(file_path)
    except Exception as e:
        print(f"Error: Could not open or parse the DOCX file at '{file_path}'.")
        print(f"Details: {e}")
        sys.exit(1)

    print("=" * 30)
    print(" DOCUMENT SECTION SETTINGS")
    print("=" * 30)

    for i, section in enumerate(document.sections):
        print(f"\n--- Settings for Section {i+1} ---")
        print(
            f"  - Page Size: {section.page_width.cm:.2f}cm (width) x {section.page_height.cm:.2f}cm (height)"
        )
        print(f"  - Orientation: {str(section.orientation)}")
        print(f"  - Margins:")
        print(f"    - Top:    {section.top_margin.cm:.2f}cm")
        print(f"    - Bottom: {section.bottom_margin.cm:.2f}cm")
        print(f"    - Left:   {section.left_margin.cm:.2f}cm")
        print(f"    - Right:  {section.right_margin.cm:.2f}cm")
        print(f"  - Header Distance: {section.header_distance.cm:.2f}cm from edge")
        print(f"  - Footer Distance: {section.footer_distance.cm:.2f}cm from edge")

    print("\n\n" + "=" * 30)
    print(" DEFINED DOCUMENT STYLES")
    print("=" * 30)

    styles = document.styles
    for style in sorted(styles, key=lambda s: s.name):
        if style.type == 1:  # Paragraph styles
            print(f"\n--- Style: '{style.name}' (Paragraph) ---")
            p_fmt = style.paragraph_format
            if p_fmt:
                print(f"  - Paragraph:")
                print(f"    - Alignment: {p_fmt.alignment}")
                line_spacing_val = (
                    f"{p_fmt.line_spacing:.2f}"
                    if p_fmt.line_spacing is not None
                    else "N/A"
                )
                print(
                    f"    - Line Spacing: {line_spacing_val} (Rule: {p_fmt.line_spacing_rule})"
                )
                space_before_pt = (
                    f"{p_fmt.space_before.pt:.2f}pt" if p_fmt.space_before else "0pt"
                )
                space_after_pt = (
                    f"{p_fmt.space_after.pt:.2f}pt" if p_fmt.space_after else "0pt"
                )
                print(
                    f"    - Spacing: {space_before_pt} before, {space_after_pt} after"
                )
                left_indent_cm = (
                    f"{p_fmt.left_indent.cm:.2f}cm" if p_fmt.left_indent else "0cm"
                )
                right_indent_cm = (
                    f"{p_fmt.right_indent.cm:.2f}cm" if p_fmt.right_indent else "0cm"
                )
                first_line_indent_cm = (
                    f"{p_fmt.first_line_indent.cm:.2f}cm"
                    if p_fmt.first_line_indent
                    else "0cm"
                )
                print(
                    f"    - Indentation: Left {left_indent_cm}, Right {right_indent_cm}, First-line {first_line_indent_cm}"
                )
                print(f"    - Widow/Orphan Control: {p_fmt.widow_control}")

            font = style.font
            if font:
                print(f"  - Font: {get_font_info(font)}")

        elif style.type == 2:  # Character styles
            print(f"\n--- Style: '{style.name}' (Character) ---")
            font = style.font
            if font:
                print(f"  - Font: {get_font_info(font)}")

    print("\n\n" + "=" * 30)
    print(" TABLE ANALYSIS")
    print("=" * 30)

    if not document.tables:
        print("No tables found in this document.")
    else:
        for i, table in enumerate(document.tables):
            print(f"\n--- Table {i+1} ---")
            print(f"  - Style: '{table.style.name}'")
            print(f"  - Alignment: {table.alignment}")
            print(f"  - Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Check if this cell is part of a vertical merge.
                    # We inspect the cell's own XML properties (_tc) to do this correctly.
                    vMerge = cell._tc.find(qn("w:vMerge"))

                    # If a vMerge tag exists but has no 'val' attribute (or the val is not 'restart'),
                    # it means this cell is merged into the cell above it. We should skip it.
                    if vMerge is not None and vMerge.get(qn("w:val")) is None:
                        continue

                    print(f"\n    - Cell({r_idx+1}, {c_idx+1})")
                    print(f"      - Vertical Align: {cell.vertical_alignment}")
                    print(f"      - Borders: {get_cell_border_info(cell)}")
                    for p_idx, p in enumerate(cell.paragraphs):
                        if p.text.strip():
                            print(
                                f"      - Paragraph {p_idx+1}: '{p.text[:60].strip()}...'"
                            )
                            for r_idx_run, run in enumerate(p.runs):
                                if run.text.strip():
                                    print(
                                        f"        - Run {r_idx_run+1}: {get_font_info(run.font)}"
                                    )

    print("\n\n" + "=" * 30)
    print(" ACTUAL FORMATTING OF DOCUMENT CONTENT (First 20 Paragraphs)")
    print("=" * 30)

    MAX_PARAGRAPHS_TO_INSPECT = 20
    for i, p in enumerate(document.paragraphs):
        if i >= MAX_PARAGRAPHS_TO_INSPECT:
            print(
                f"\n[... Stopped inspection after {MAX_PARAGRAPHS_TO_INSPECT} paragraphs ...]"
            )
            break

        if not p.text.strip():
            continue

        print(f"\n--- Paragraph {i+1} ---")
        print(f"  - Style: '{p.style.name}'")
        print(f"  - Text: '{p.text[:80]}...'")

        # Check formatting for each run in the paragraph
        for j, run in enumerate(p.runs):
            if not run.text.strip():
                continue
            print(f"    - Run {j+1}: {get_font_info(run.font)} | Text: '{run.text}'")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Extract formatting and style information from a .docx file.",
        formatter_class=argparse.RawTextHelpFormatter,
    )
    parser.add_argument(
        "docx_path",
        type=str,
        help="The full path to the .docx file you want to inspect.\nExample: python style_inspector.py '/Users/yourname/Documents/template.docx'",
    )
    parser.add_argument(
        "-o",
        "--output",
        type=str,
        default=None,
        help="Path to save the output report as a .txt file.\nIf not provided, output is printed to the console.",
    )
    args = parser.parse_args()

    if args.output:
        original_stdout = sys.stdout
        print(f"Starting inspection. Output will be saved to: {args.output}")
        try:
            with open(args.output, "w", encoding="utf-8") as f:
                sys.stdout = f
                inspect_docx_styles(args.docx_path)
                print("\n\nInspection complete.")
        except Exception as e:
            # Still print error to console if file writing fails
            original_stdout.write(f"An error occurred: {e}\n")
        finally:
            sys.stdout = original_stdout  # Restore standard output
        print("Inspection finished.")
    else:
        inspect_docx_styles(args.docx_path)
        print("\n\nInspection complete.")
